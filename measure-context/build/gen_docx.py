# -*- coding: utf-8 -*-
"""Word (.docx) renderer with python-docx.

Body -> headings + paragraphs, tables -> native Word tables, the two captures ->
embedded pictures (add_picture = real image bytes). Diagrams are rendered as
native Word relationship tables (python-docx has no autoshape/connector API);
this limitation is reported, not worked around with fragile XML injection.

build() returns (path, logical_text) where logical_text is every string written
to a heading / paragraph / table cell -- the (a) logical-content measure.
"""
import os
from docx import Document
from docx.shared import Inches, Pt

import content as C
import figdata as F
from gen_text import SAMPLES, _meta_lines


def _add_table(doc, header, rows, L):
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Table Grid"
    for c, h in enumerate(header):
        cell = t.rows[0].cells[c]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        L.append(str(h))
    for ri, row in enumerate(rows):
        for c, val in enumerate(row):
            t.rows[ri + 1].cells[c].text = str(val)
            L.append(str(val))


def _caption(doc, text, L):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    L.append(text)


def _rel_table(doc, spec, L):
    title, header, rows = spec
    _caption(doc, title, L)
    _add_table(doc, header, rows, L)


def build():
    doc = Document()
    L = []

    doc.add_heading(C.DOC["title"], level=0)
    L.append(C.DOC["title"])
    meta = " / ".join("%s: %s" % kv for kv in _meta_lines())
    doc.add_paragraph(meta)
    L.append(meta)
    doc.add_paragraph("概要: " + C.DOC["abstract"])
    L.append("概要: " + C.DOC["abstract"])

    dt = F.diagram_tables()

    def section(num):
        n, head, paras = next(s for s in C.SECTIONS if s[0] == num)
        doc.add_heading("%s %s" % (n, head), level=1)
        L.append("%s %s" % (n, head))
        for p in paras:
            doc.add_paragraph(p)
            L.append(p)

    section("1")
    for spec in dt["block"]:
        _rel_table(doc, spec, L)
    section("2")
    section("3")
    for spec in dt["state"]:
        _rel_table(doc, spec, L)
    section("4")
    for spec in dt["sequence"]:
        _rel_table(doc, spec, L)
    section("5")
    for spec in dt["flow"]:
        _rel_table(doc, spec, L)
    section("6")
    section("7")
    _caption(doc, C.PARAM_TABLE_CAPTION, L)
    _add_table(doc, C.PARAM_TABLE_HEADER, C.PARAM_TABLE_ROWS, L)
    _caption(doc, C.IF_TABLE_CAPTION, L)
    _add_table(doc, C.IF_TABLE_HEADER, C.IF_TABLE_ROWS, L)

    for cap, img in [(C.FIG_CAP1, C.CAPTURE1), (C.FIG_CAP2, C.CAPTURE2)]:
        _caption(doc, cap, L)
        doc.add_picture(os.path.join(SAMPLES, img), width=Inches(5.8))

    path = os.path.join(SAMPLES, "spec_lkas.docx")
    doc.save(path)
    return path, "\n".join(L)


if __name__ == "__main__":
    from common import count_tokens
    p, logical = build()
    print("docx logical:", count_tokens(logical), "tokens,", len(logical), "chars ->", os.path.basename(p))
